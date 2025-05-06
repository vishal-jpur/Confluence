from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_doc(filename):
    doc = Document()

    # Title
    title = doc.add_heading('Monthly Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Prepared by: Your Name\nDate: 06-May-2025')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Add a blank line

    # Section Heading
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(
        'This report summarizes the activities and progress for the month. The major highlights are listed below.'
    )

    # Bullet Points
    points = [
        'Completed the performance testing framework.',
        'Deployed the automation script to CI/CD pipeline.',
        'Reduced average response time by 20%.'
    ]
    for point in points:
        doc.add_paragraph(point, style='List Bullet')

    # Add a Table
    doc.add_heading('Performance Metrics', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Page Name'
    hdr_cells[1].text = 'Avg Response Time (ms)'
    hdr_cells[2].text = 'Error Rate (%)'

    data = [
        ('Login', '1500', '1.2'),
        ('Dashboard', '1800', '0.5'),
        ('Reports', '1900', '0.9'),
    ]

    for row in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = value

    # Save the Document
    doc.save(filename)
    print(f'Document saved as {filename}')

# Usage
create_word_doc('Monthly_Report.docx')
